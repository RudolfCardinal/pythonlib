#!/usr/bin/python

"""
Converts a bunch of stuff to text, either from external files or from in-memory
binary objects (BLOBs).

Prerequisites:

    sudo apt-get install antiword
    pip install docx pdfminer

Author: Rudolf Cardinal (rudolf@pobox.com)
Created: Feb 2015
Last update: 24 Sep 2015

Copyright/licensing:

    Copyright (C) 2015-2015 Rudolf Cardinal (rudolf@pobox.com).

    Licensed under the Apache License, Version 2.0 (the "License");
    you may not use this file except in compliance with the License.
    You may obtain a copy of the License at

        http://www.apache.org/licenses/LICENSE-2.0

    Unless required by applicable law or agreed to in writing, software
    distributed under the License is distributed on an "AS IS" BASIS,
    WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
    See the License for the specific language governing permissions and
    limitations under the License.

See also:
    Word
        http://stackoverflow.com/questions/125222
        http://stackoverflow.com/questions/42482
    PDF
        http://stackoverflow.com/questions/25665
        https://pypi.python.org/pypi/slate
        http://stackoverflow.com/questions/5725278
    RTF
        unrtf
        http://superuser.com/questions/243084/rtf-to-txt-on-unix
    Multi-purpose:
        https://pypi.python.org/pypi/fulltext/
        https://media.readthedocs.org/pdf/textract/latest/textract.pdf
    DOCX
        http://etienned.github.io/posts/extract-text-from-word-docx-simply/
"""


# =============================================================================
# Imports
# =============================================================================

from __future__ import division, print_function, absolute_import
import argparse
try:  # http://python3porting.com/stdlib.html
    from io import StringIO  # Python 3
except ImportError:
    # noinspection PyCompatibility,PyUnresolvedReferences
    from cStringIO import StringIO  # Python 2
import io
import os
import re
import shutil
import subprocess
import sys
import textwrap
from xml.etree import ElementTree as ElementTree
# ... cElementTree used to be the fast implementation; now ElementTree is fast
# and cElementTree is deprecated; see
# https://docs.python.org/3.4/library/xml.etree.elementtree.html
import zipfile

# noinspection PyPackageRequirements
import bs4  # pip install beautifulsoup4
try:
    # noinspection PyPackageRequirements
    import docx  # pip install python-docx (NOT docx) - BUT python-docx requires lxml which has C dependencies  # noqa
    import docx.document
    import docx.oxml.table
    import docx.oxml.text.paragraph
    import docx.table
    import docx.text.paragraph
except ImportError:
    docx = None
try:
    import docx2txt  # pip install docx2txt
except ImportError:
    docx2txt = None
try:
    # noinspection PyPackageRequirements
    import pdfminer  # pip install pdfminer
    # noinspection PyPackageRequirements
    import pdfminer.pdfinterp
    # noinspection PyPackageRequirements
    import pdfminer.converter
    # noinspection PyPackageRequirements
    import pdfminer.layout
    # noinspection PyPackageRequirements
    import pdfminer.pdfpage
except ImportError:
    pdfminer = None
# noinspection PyPackageRequirements
import prettytable  # pip install PrettyTable
try:
    # noinspection PyPackageRequirements
    import pyth  # pip install pyth (PYTHON 2 ONLY; https://pypi.python.org/pypi/pyth/0.5.4)  # noqa
    # noinspection PyPackageRequirements
    import pyth.plugins.rtf15.reader
    # noinspection PyPackageRequirements
    import pyth.plugins.plaintext.writer
except ImportError:
    pyth = None
import six
# import texttable  # ... can't deal with Unicode properly

import logging
log = logging.getLogger(__name__)
log.addHandler(logging.NullHandler())

# =============================================================================
# Constants
# =============================================================================

ENCODING = "utf-8"

# =============================================================================
# External tool map
# =============================================================================

if six.PY2:
    tools = {
        'antiword': 'antiword',
        'pdftotext': 'pdftotext',
        'strings': 'strings',
        'strings2': 'strings2',
        'unrtf': 'unrtf',
    }
else:
    tools = {
        'antiword': shutil.which('antiword'),  # sudo apt-get install antiword
        'pdftotext': shutil.which('pdftotext'),  # core part of Linux?
        'strings': shutil.which('strings'),  # part of standard Unix
        'strings2': shutil.which('strings2'),
        # ... Windows: https://technet.microsoft.com/en-us/sysinternals/strings.aspx  # noqa
        # ... Windows: http://split-code.com/strings2.html
        'unrtf': shutil.which('unrtf'),  # sudo apt-get install unrtf
    }


def update_external_tools(tooldict):
    global tools
    tools.update(tooldict)


# =============================================================================
# Support functions
# =============================================================================

def get_filelikeobject(filename=None, blob=None):
    """Guard the use of this function with 'with'."""
    if not filename and not blob:
        raise ValueError("no filename and no blob")
    if filename and blob:
        raise ValueError("specify either filename or blob")
    if filename:
        return open(filename, 'rb')
    else:
        return io.BytesIO(blob)


def get_file_contents(filename=None, blob=None):
    """Returns binary contents of a file, or blob."""
    if not filename and not blob:
        raise ValueError("no filename and no blob")
    if filename and blob:
        raise ValueError("specify either filename or blob")
    if blob:
        return blob
    with open(filename, 'rb') as f:
        return f.read()


def get_cmd_output(*args, **kwargs):
    """Returns text output of a command."""
    encoding = kwargs.get("encoding", ENCODING)
    log.debug("get_cmd_output(): args = {}".format(repr(args)))
    p = subprocess.Popen(args, stdout=subprocess.PIPE)
    stdout, stderr = p.communicate()
    return stdout.decode(encoding, errors='ignore')


def get_cmd_output_from_stdin(stdint_content_binary, *args, **kwargs):
    """Returns text output of a command, passing binary data in via stdin."""
    encoding = kwargs.get("encoding", ENCODING)
    p = subprocess.Popen(args, stdin=subprocess.PIPE, stdout=subprocess.PIPE)
    stdout, stderr = p.communicate(input=stdint_content_binary)
    return stdout.decode(encoding, errors='ignore')


# =============================================================================
# PDF
# =============================================================================

# noinspection PyUnresolvedReferences
def convert_pdf_to_txt(filename=None, blob=None):
    """Pass either a filename or a binary object."""
    pdftotext = tools['pdftotext']
    if pdftotext:  # External command method
        if filename:
            return get_cmd_output(pdftotext, filename, '-')
        else:
            return get_cmd_output_from_stdin(blob, pdftotext, '-', '-')
    elif pdfminer:  # Memory-hogging method
        with get_filelikeobject(filename, blob) as fp:
            rsrcmgr = pdfminer.pdfinterp.PDFResourceManager()
            retstr = StringIO()
            codec = ENCODING
            laparams = pdfminer.layout.LAParams()
            device = pdfminer.converter.TextConverter(
                rsrcmgr, retstr, codec=codec, laparams=laparams)
            interpreter = pdfminer.pdfinterp.PDFPageInterpreter(rsrcmgr,
                                                                device)
            password = ""
            maxpages = 0
            caching = True
            pagenos = set()
            for page in pdfminer.pdfpage.PDFPage.get_pages(
                    fp, pagenos, maxpages=maxpages, password=password,
                    caching=caching, check_extractable=True):
                interpreter.process_page(page)
            text = retstr.getvalue().decode(ENCODING)
        return text
    else:
        raise AssertionError("No PDF-reading tool available")


def availability_pdf():
    pdftotext = tools['pdftotext']
    if pdftotext:
        return True
    elif pdfminer:
        log.warning("PDF conversion: pdftotext missing; "
                    "using pdfminer (less efficient)")
        return True
    else:
        return False


# =============================================================================
# DOCX
# =============================================================================

# -----------------------------------------------------------------------------
# Generic
# -----------------------------------------------------------------------------

def docx_process_simple_text(text, width):
    if width:
        return '\n'.join(textwrap.wrap(text, width=width))
    else:
        return text


def docx_process_table(table, width, min_col_width):
    """
    Structure:
        table
            .rows[]
                .cells[]
                    .paragraphs[]
                        .text
    That's the structure of a docx.table.Table object, but also of our homebrew
    creation, CustomDocxTable.
    """
    ncols = 1
    for row in table.rows:
        ncols = max(ncols, len(row.cells))
    pt = prettytable.PrettyTable(
        field_names=list(range(ncols)),
        encoding=ENCODING,
        header=False,
        border=True,
        hrules=prettytable.ALL,
        vrules=prettytable.ALL,
    )
    pt.align = 'l'
    pt.valign = 't'
    pt.max_width = max(width // ncols, min_col_width)
    for row in table.rows:
        ncols = max(ncols, len(row.cells))
        ptrow = []
        for cell in row.cells:
            cellparagraphs = [paragraph.text.strip()
                              for paragraph in cell.paragraphs]
            cellparagraphs = [x for x in cellparagraphs if x]
            ptrow.append('\n\n'.join(cellparagraphs))
        pt.add_row(ptrow)
    return pt.get_string()


# -----------------------------------------------------------------------------
# With the docx library
# -----------------------------------------------------------------------------


# noinspection PyProtectedMember
def docx_docx_iter_block_items(parent):
    """
    https://github.com/python-openxml/python-docx/issues/40

    Yield each paragraph and table child within *parent*, in document order.
    Each returned value is an instance of either Table or Paragraph. *parent*
    would most commonly be a reference to a main Document object, but
    also works for a _Cell object, which itself can contain paragraphs and
    tables.

    NOTE: uses internals of the python-docx (docx) library; subject to change;
    this version works with docx 0.8.5
    """
    if isinstance(parent, docx.document.Document):
        parent_elm = parent.element.body
    elif isinstance(parent, docx.table._Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("something's not right")

    for child in parent_elm.iterchildren():
        if isinstance(child, docx.oxml.text.paragraph.CT_P):
            yield docx.text.paragraph.Paragraph(child, parent)
        elif isinstance(child, docx.oxml.table.CT_Tbl):
            yield docx.table.Table(child, parent)


def docx_docx_gen_text(doc, width, min_col_width, in_order=True):
    if in_order:
        for thing in docx_docx_iter_block_items(doc):
            if isinstance(thing, docx.text.paragraph.Paragraph):
                yield docx_process_simple_text(thing.text, width)
            elif isinstance(thing, docx.table.Table):
                yield docx_process_table(thing, width, min_col_width)
    else:
        for paragraph in doc.paragraphs:
            yield docx_process_simple_text(paragraph.text, width)
        for table in doc.tables:
            yield docx_process_table(table, width, min_col_width)


# -----------------------------------------------------------------------------
# In a D.I.Y. fashion
# -----------------------------------------------------------------------------
# DOCX specification: http://www.ecma-international.org/news/TC45_current_work/TC45_available_docs.htm  # noqa

DOCX_HEADER_FILE_REGEX = re.compile('word/header[0-9]*.xml')
DOCX_DOC_FILE = 'word/document.xml'
DOCX_FOOTER_FILE_REGEX = re.compile('word/footer[0-9]*.xml')
DOCX_SCHEMA_URL = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'  # noqa


def docx_qn(tagroot):
    return '{{{}}}{}'.format(DOCX_SCHEMA_URL, tagroot)


DOCX_TEXT = docx_qn('t')
DOCX_TABLE = docx_qn('tbl')  # https://github.com/python-openxml/python-docx/blob/master/docx/table.py  # noqa
DOCX_TAB = docx_qn('tab')
DOCX_NEWLINES = [docx_qn('br'), docx_qn('cr')]
DOCX_NEWPARA = docx_qn('p')
DOCX_TABLE_ROW = docx_qn('tr')
DOCX_TABLE_CELL = docx_qn('tc')


def gen_xml_files_from_docx(fp):
    z = zipfile.ZipFile(fp)
    filelist = z.namelist()
    for filename in filelist:
        if DOCX_HEADER_FILE_REGEX.match(filename):
            yield z.read(filename)
    yield z.read(DOCX_DOC_FILE)
    for filename in filelist:
        if DOCX_FOOTER_FILE_REGEX.match(filename):
            yield z.read(filename)


def docx_text_from_xml(xml, **kwargs):
    root = ElementTree.fromstring(xml)
    return docx_text_from_xml_node(root, 0, **kwargs)


def docx_text_from_xml_node(node, level, **kwargs):
    text = ''
    # log.debug("Level {}, tag {}".format(level, node.tag))
    if node.tag == DOCX_TEXT:
        text += node.text or ''
    elif node.tag == DOCX_TAB:
        text += '\t'
    elif node.tag in DOCX_NEWLINES:
        text += '\n'
    elif node.tag == DOCX_NEWPARA:
        text += '\n\n'

    if node.tag == DOCX_TABLE:
        text += '\n\n' + docx_table_from_xml_node(node, level, **kwargs)
    else:
        for child in node:
            text += docx_text_from_xml_node(child, level + 1, **kwargs)
    return text


class CustomDocxTable(object):
    def __init__(self, rows=None):
        self.rows = rows or []

    def add_row(self, row):
        self.rows.append(row)

    def new_row(self):
        self.rows.append(CustomDocxTableRow())

    def new_cell(self):
        self.rows[-1].new_cell()

    def add_paragraph(self, text):
        self.rows[-1].add_paragraph(text)

    def __repr__(self):
        return "CustomDocxTable(rows={})".format(repr(self.rows))


class CustomDocxTableRow(object):
    def __init__(self, cells=None):
        self.cells = cells or []

    def add_cell(self, cell):
        self.cells.append(cell)

    def new_cell(self):
        self.cells.append(CustomDocxTableCell())

    def add_paragraph(self, text):
        self.cells[-1].add_paragraph(text)

    def __repr__(self):
        return "CustomDocxTableRow(cells={})".format(repr(self.cells))


class CustomDocxTableCell(object):
    def __init__(self, paragraphs=None):
        self.paragraphs = paragraphs or []

    def add_paragraph(self, text):
        self.paragraphs.append(CustomDocxParagraph(text))

    def __repr__(self):
        return "CustomDocxTableCell(paragraphs={})".format(
            repr(self.paragraphs))


class CustomDocxParagraph(object):
    def __init__(self, text=''):
        self.text = text or ''

    def __repr__(self):
        return "CustomDocxParagraph(text={})".format(repr(self.text))


def docx_table_from_xml_node(table_node, level, **kwargs):
    table = CustomDocxTable()
    for row_node in table_node:
        if row_node.tag != DOCX_TABLE_ROW:
            continue
        table.new_row()
        for cell_node in row_node:
            if cell_node.tag != DOCX_TABLE_CELL:
                continue
            table.new_cell()
            for para_node in cell_node:
                text = docx_text_from_xml_node(para_node, level, **kwargs)
                if text:
                    table.add_paragraph(text)
    return docx_process_table(table, **kwargs)


def convert_docx_to_text(filename=None, blob=None,
                         width=120, min_col_width=15):
    """
    Pass either a filename or a binary object.

    -   Old docx (https://pypi.python.org/pypi/python-docx) has been superseded
        (see https://github.com/mikemaccana/python-docx).
        -   docx.opendocx(file) uses zipfile.ZipFile, which can take either a
            filename or a file-like object
            (https://docs.python.org/2/library/zipfile.html).
        -   Method was:
                with get_filelikeobject(filename, blob) as fp:
                    document = docx.opendocx(fp)
                    paratextlist = docx.getdocumenttext(document)
                return '\n\n'.join(paratextlist)

    -   Newer docx is python-docx
            https://pypi.python.org/pypi/python-docx
            https://python-docx.readthedocs.org/en/latest/
            http://stackoverflow.com/questions/25228106
        However, it uses lxml, which has C dependencies, so it doesn't always
        install properly on e.g. bare Windows machines.

        PERFORMANCE of my method:
            -   nice table formatting
            -   but tables grouped at end, not in sensible places
            -   can iterate via "doc.paragraphs" and "doc.tables" but not in
                true document order, it seems
            -   others have noted this too:
                https://github.com/python-openxml/python-docx/issues/40
                https://github.com/deanmalmgren/textract/pull/92

    -   docx2txt is at https://pypi.python.org/pypi/docx2txt/0.6; this is
        pure Python. Its command-line function appears to be for Python 2 only
        (2016-04-21: crashes under Python 3; is due to an encoding bug).
        However, it seems fine as a library.
        It doesn't handle in-memory blobs properly, though, so we need to
        extend it.

        PERFORMANCE OF ITS process() function:
        - all text comes out
        - table text is in a sensible place
        - table formatting is lost.

    -   Other manual methods (not yet implemented):
        http://etienned.github.io/posts/extract-text-from-word-docx-simply/

        ... looks like it won't deal with header stuff (etc.) that docx2txt
            handles.

    -   Upshot: we need a DIY version.

    -   See also this "compile lots of techniques" libraries, which has C
        dependencies:
            http://textract.readthedocs.org/en/latest/

    """

    if True:
        text = ''
        with get_filelikeobject(filename, blob) as fp:
            for xml in gen_xml_files_from_docx(fp):
                text += docx_text_from_xml(xml, width=width,
                                           min_col_width=min_col_width)
        return text
    elif docx:
        with get_filelikeobject(filename, blob) as fp:
            document = docx.Document(fp)
            return '\n\n'.join(
                docx_docx_gen_text(document, width=width,
                                   min_col_width=min_col_width))
    elif docx2txt:
        if filename:
            return docx2txt.process(filename)
        else:
            raise NotImplementedError("docx2txt BLOB handling not written")
    else:
        raise AssertionError("No DOCX-reading tool available")


# =============================================================================
# ODT
# =============================================================================

def convert_odt_to_text(filename=None, blob=None):
    """Pass either a filename or a binary object."""
    # We can't use exactly the same method as for DOCX files, using docx:
    # sometimes that works, but sometimes it falls over with:
    # KeyError: "There is no item named 'word/document.xml' in the archive"
    with get_filelikeobject(filename, blob) as fp:
        z = zipfile.ZipFile(fp)
        tree = ElementTree.fromstring(z.read('content.xml'))
        # ... may raise zipfile.BadZipfile
        textlist = []
        for element in tree.iter():
            if element.text:
                textlist.append(element.text.strip())
    return '\n\n'.join(textlist)


# =============================================================================
# HTML
# =============================================================================

def convert_html_to_text(filename=None, blob=None):
    with get_filelikeobject(filename, blob) as fp:
        soup = bs4.BeautifulSoup(fp)
        return soup.get_text()


# =============================================================================
# XML
# =============================================================================


def convert_xml_to_text(filename=None, blob=None):
    with get_filelikeobject(filename, blob) as fp:
        soup = bs4.BeautifulStoneSoup(fp)
        return soup.get_text()


# =============================================================================
# RTF
# =============================================================================


# noinspection PyUnresolvedReferences
def convert_rtf_to_text(filename=None, blob=None):
    unrtf = tools['unrtf']
    if unrtf:  # Best
        if filename:
            return get_cmd_output(
                unrtf, '--text', '--nopict', '--quiet', filename)
        else:
            return get_cmd_output_from_stdin(
                blob, unrtf, '--text', '--nopict', '--quiet')
    elif pyth:  # Very memory-consuming:
        # https://github.com/brendonh/pyth/blob/master/pyth/plugins/rtf15/reader.py  # noqa
        with get_filelikeobject(filename, blob) as fp:
            doc = pyth.plugins.rtf15.reader.Rtf15Reader.read(fp)
        return (
            pyth.plugins.plaintext.writer.PlaintextWriter.write(doc).getvalue()
        )
    else:
        raise AssertionError("No RTF-reading tool available")


def availability_rtf():
    unrtf = tools['unrtf']
    if unrtf:
        return True
    elif pyth:
        log.warning("RTF conversion: unrtf missing; "
                    "using pyth (less efficient)")
        return True
    else:
        return False


# =============================================================================
# DOC
# =============================================================================

def convert_doc_to_text(filename=None, blob=None):
    antiword = tools['antiword']
    if antiword:
        if filename:
            return get_cmd_output(antiword, filename)
        else:
            return get_cmd_output_from_stdin(blob, antiword, '-')
    else:
        raise AssertionError("No DOC-reading tool available")


def availability_doc():
    antiword = tools['antiword']
    return bool(antiword)


# =============================================================================
# Anything
# =============================================================================

def convert_anything_to_text(filename=None, blob=None):
    # strings is a standard Unix command to get text from any old rubbish
    strings = tools['strings'] or tools['strings2']
    if strings:
        if filename:
            return get_cmd_output(strings, filename)
        else:
            return get_cmd_output_from_stdin(blob, strings)
    else:
        raise AssertionError("No fallback string-reading tool available")


def availability_anything():
    strings = tools['strings'] or tools['strings2']
    return bool(strings)


# =============================================================================
# Decider
# =============================================================================

ext_map = {
    # Converter functions must be of the form func(filename, blob):
    # Availability must be either a boolean or a function that takes no params.
    '.doc': {
        'converter': convert_doc_to_text,
        'availability': availability_doc,
    },
    '.dot': {
        'converter': convert_doc_to_text,
        'availability': availability_doc,
    },
    '.docm': {
        'converter': convert_docx_to_text,
        'availability': True,
    },
    '.docx': {
        'converter': convert_docx_to_text,
        'availability': True,
    },
    '.html': {
        'converter': convert_html_to_text,
        'availability': True,
    },
    '.htm': {
        'converter': convert_html_to_text,
        'availability': True,
    },
    '.log': {
        'converter': get_file_contents,
        'availability': True,
    },
    '.odt': {
        'converter': convert_odt_to_text,
        'availability': True,
    },
    '.pdf': {
        'converter': convert_pdf_to_txt,
        'availability': availability_pdf,
    },
    '.rtf': {
        'converter': convert_rtf_to_text,
        'availability': availability_rtf,
    },
    '.xml': {
        'converter': convert_xml_to_text,
        'availability': True,
    },
    '.txt': {
        'converter': get_file_contents,
        'availability': True,
    },
    None: {  # fallback
        'converter': convert_anything_to_text,
        'availability': availability_anything,
    },
}


def document_to_text(filename=None, blob=None, extension=None):
    """Pass either a filename or a binary object.
    - Raises an exception for malformed arguments, missing files, bad
      filetypes, etc.
    - Returns a string if the file was processed (potentially an empty string).
    """
    if not filename and not blob:
        raise ValueError("document_to_text: no filename and no blob")
    if filename and blob:
        raise ValueError("document_to_text: specify either filename or blob")
    if blob and not extension:
        raise ValueError("document_to_text: need extension hint for blob")
    if filename:
        stub, extension = os.path.splitext(filename)
    else:
        if extension[0] != ".":
            extension = "." + extension
    extension = extension.lower()

    # Ensure blob is an appropriate type
    log.debug(
        "filename: {}, blob type: {}, blob length: {}, extension: {}".format(
            filename,
            type(blob),
            len(blob) if blob is not None else None,
            extension))

    # Choose method
    info = ext_map.get(extension)
    if info is None:
        log.warning("Unknown filetype: {}; using generic tool".format(
            extension))
        info = ext_map[None]
    func = info['converter']
    return func(filename, blob)


def is_text_extractor_available(extension):
    if extension is not None:
        extension = extension.lower()
    info = ext_map.get(extension)
    if info is None:
        return False
    availability = info['availability']
    if type(availability) == bool:
        return availability
    elif callable(availability):
        return availability()
    else:
        raise ValueError(
            "Bad information object for extension: {}".format(extension))


def require_text_extractor(extension):
    if not is_text_extractor_available(extension):
        raise ValueError(
            "No text extractor available for extension: {}".format(extension))


# =============================================================================
# main, for command-line use
# =============================================================================

def main():
    logging.basicConfig(level=logging.DEBUG)
    parser = argparse.ArgumentParser()
    parser.add_argument("inputfile", nargs="?", help="Input file name")
    parser.add_argument(
        "--availability", nargs='*',
        help="File extensions to check availability for (use a '.' prefix, "
             "and use the special extension 'None' to check the fallback "
             "processor")
    args = parser.parse_args()
    if args.availability:
        for ext in args.availability:
            if ext.lower() == 'none':
                ext = None
            available = is_text_extractor_available(ext)
            print("Extractor for extension {} present: {}".format(ext,
                                                                  available))
        return
    if not args.inputfile:
        parser.print_help(sys.stderr)
        return
    result = document_to_text(filename=args.inputfile)
    if result is None:
        return
    elif six.PY2 and isinstance(result, six.text_type):
        print(result.encode(ENCODING))
    else:
        print(result)


if __name__ == '__main__':
    main()
