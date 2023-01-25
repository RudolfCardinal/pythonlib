#!/usr/bin/env python
# cardinal_pythonlib/extract_text.py

"""
===============================================================================

    Original code copyright (C) 2009-2022 Rudolf Cardinal (rudolf@pobox.com).

    This file is part of cardinal_pythonlib.

    Licensed under the Apache License, Version 2.0 (the "License");
    you may not use this file except in compliance with the License.
    You may obtain a copy of the License at

        https://www.apache.org/licenses/LICENSE-2.0

    Unless required by applicable law or agreed to in writing, software
    distributed under the License is distributed on an "AS IS" BASIS,
    WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
    See the License for the specific language governing permissions and
    limitations under the License.

===============================================================================

**Converts a bunch of stuff to text, either from external files or from
in-memory binary objects (BLOBs).**

Prerequisites:

.. code-block:: bash

    sudo apt-get install antiword
    pip install docx pdfminer

- Author: Rudolf Cardinal (rudolf@pobox.com)
- Created: Feb 2015
- Last update: 24 Sep 2015

See also:

- Word

  - https://stackoverflow.com/questions/125222
  - https://stackoverflow.com/questions/424822

- PDF

  - https://stackoverflow.com/questions/25665
  - https://pypi.python.org/pypi/slate
  - https://stackoverflow.com/questions/5725278

- RTF

  - unrtf
  - https://superuser.com/questions/243084/rtf-to-txt-on-unix

- Multi-purpose:

  - https://pypi.python.org/pypi/fulltext/
  - https://media.readthedocs.org/pdf/textract/latest/textract.pdf

- DOCX

  - https://etienned.github.io/posts/extract-text-from-word-docx-simply/

"""


# =============================================================================
# Imports
# =============================================================================

import argparse
from io import StringIO  # Python 3
import io
import logging
import os
import re
import shutil
import subprocess
import sys
import textwrap
from typing import (
    BinaryIO,
    Dict,
    Generator,
    Iterable,
    Iterator,
    List,
    Optional,
    Union,
)
from xml.etree import ElementTree as ElementTree

# ... cElementTree used to be the fast implementation; now ElementTree is fast
# and cElementTree is deprecated; see
# https://docs.python.org/3.4/library/xml.etree.elementtree.html
import zipfile

import bs4
import prettytable
from semantic_version import Version

# import texttable  # ... can't deal with Unicode properly

from cardinal_pythonlib.logs import get_brace_style_log_with_null_handler

try:
    import chardet
    from chardet.universaldetector import UniversalDetector
except ImportError:
    chardet = None
    UniversalDetector = None

try:
    # noinspection PyPackageRequirements
    import docx  # pip install python-docx (NOT docx) - BUT python-docx requires lxml which has C dependencies  # noqa

    # noinspection PyPackageRequirements
    import docx.document

    # noinspection PyPackageRequirements
    import docx.oxml.table

    # noinspection PyPackageRequirements
    import docx.oxml.text.paragraph

    # noinspection PyPackageRequirements
    import docx.table

    # noinspection PyPackageRequirements
    import docx.text.paragraph

    DOCX_DOCUMENT_TYPE = "docx.document.Document"
    DOCX_TABLE_TYPE = Union["docx.table.Table", "CustomDocxTable"]
    DOCX_CONTAINER_TYPE = Union[DOCX_DOCUMENT_TYPE, "docx.table._Cell"]
    DOCX_BLOCK_ITEM_TYPE = Union[
        "docx.text.paragraph.Paragraph", "docx.table.Table"
    ]
except ImportError:
    docx = None
    DOCX_DOCUMENT_TYPE = None
    DOCX_TABLE_TYPE = "CustomDocxTable"
    DOCX_CONTAINER_TYPE = None
    DOCX_BLOCK_ITEM_TYPE = None

try:
    import docx2txt  # pip install docx2txt
except ImportError:
    docx2txt = None

try:
    # noinspection PyPackageRequirements
    import pdfminer  # pip install pdfminer

    # noinspection PyPackageRequirements
    import pdfminer.pdfinterp

    # noinspection PyPackageRequirements
    import pdfminer.converter

    # noinspection PyPackageRequirements
    import pdfminer.layout

    # noinspection PyPackageRequirements
    import pdfminer.pdfpage
except ImportError:
    pdfminer = None

try:
    # noinspection PyPackageRequirements
    import pyth  # pip install pyth (PYTHON 2 ONLY; https://pypi.python.org/pypi/pyth/0.5.4)  # noqa

    # noinspection PyPackageRequirements
    import pyth.plugins.rtf15.reader

    # noinspection PyPackageRequirements
    import pyth.plugins.plaintext.writer
except ImportError:
    pyth = None

log = get_brace_style_log_with_null_handler(__name__)

# =============================================================================
# Constants
# =============================================================================

AVAILABILITY = "availability"
CONVERTER = "converter"
DEFAULT_WIDTH = 120
DEFAULT_MIN_COL_WIDTH = 15
SYS_ENCODING = sys.getdefaultencoding()
ENCODING = "utf-8"

# =============================================================================
# External tool map
# =============================================================================

tools = {
    "antiword": shutil.which("antiword"),  # sudo apt-get install antiword
    "pdftotext": shutil.which("pdftotext"),  # core part of Linux?
    "strings": shutil.which("strings"),  # part of standard Unix
    "strings2": shutil.which("strings2"),
    # ... Windows: https://technet.microsoft.com/en-us/sysinternals/strings.aspx  # noqa
    # ... Windows: http://split-code.com/strings2.html
    "unrtf": shutil.which("unrtf"),  # sudo apt-get install unrtf
}


def does_unrtf_support_quiet() -> bool:
    """
    The unrtf tool supports the '--quiet' argument from a version that I'm not
    quite sure of, where ``0.19.3 < version <= 0.21.9``. We check against
    0.21.9 here.
    """
    required_unrtf_version = Version("0.21.9")
    # ... probably: http://hg.savannah.gnu.org/hgweb/unrtf/
    # ... 0.21.9 definitely supports --quiet
    # ... 0.19.3 definitely doesn't support it
    unrtf_filename = shutil.which("unrtf")
    if not unrtf_filename:
        return False
    p = subprocess.Popen(
        ["unrtf", "--version"], stdout=subprocess.PIPE, stderr=subprocess.PIPE
    )
    _, err_bytes = p.communicate()
    text = err_bytes.decode(sys.getdefaultencoding())
    lines = text.split()
    if len(lines) < 1:
        return False
    version_str = lines[0]
    unrtf_version = Version(version_str)
    return unrtf_version >= required_unrtf_version


UNRTF_SUPPORTS_QUIET = does_unrtf_support_quiet()


def update_external_tools(tooldict: Dict[str, str]) -> None:
    """
    Update the global map of tools.

    Args:
        tooldict: dictionary whose keys are tools names and whose values are
            paths to the executables
    """
    global tools
    tools.update(tooldict)


# =============================================================================
# Text-processing config class
# =============================================================================


class TextProcessingConfig(object):
    """
    Class to manage control parameters for text extraction, without having
    to pass a lot of mysterious ``**kwargs`` around and lose track of what it
    means.

    All converter functions take one of these objects as a parameter.
    """

    def __init__(
        self,
        encoding: str = None,
        width: int = DEFAULT_WIDTH,
        min_col_width: int = DEFAULT_MIN_COL_WIDTH,
        plain: bool = False,
        semiplain: bool = False,
        docx_in_order: bool = True,
        horizontal_char="─",
        vertical_char="│",
        junction_char="┼",
        plain_table_start: str = None,
        plain_table_end: str = None,
        plain_table_col_boundary: str = None,
        plain_table_row_boundary: str = None,
        rstrip: bool = True,
    ) -> None:
        """
        Args:
            encoding:
                optional text file encoding to try in addition to
                :func:`sys.getdefaultencoding`.
            width:
                overall word-wrapping width
            min_col_width:
                minimum column width for tables
            plain:
                as plain as possible (e.g. for natural language processing);
                see :func:`docx_process_table`.
            semiplain:
                quite plain, but with some ASCII art representation of the
                table structure.
            docx_in_order:
                for DOCX files: if ``True``, process paragraphs and tables in
                the order they occur; if ``False``, process all paragraphs
                followed by all tables
            rstrip:
                Right-strip whitespace from all lines?
            horizontal_char:
                horizontal character to use with PrettyTable, e.g. ``-`` or
                ``─``
            vertical_char:
                vertical character to use with PrettyTable, e.g. ``|`` or
                ``│``
            junction_char:
                junction character to use with PrettyTable, e.g. ``+`` or
                ``┼``
            plain_table_start:
                table start line to use with ``plain=True``
            plain_table_end:
                table end line to use with ``plain=True``
            plain_table_col_boundary:
                boundary between columns to use with ``plain==True``
            plain_table_row_boundary:
                boundary between rows to use with ``plain==True``

        Example of a DOCX table processed with:

        - ``plain=False, semiplain=False``

          .. code-block:: none

            ┼─────────────┼─────────────┼
            │ Row 1 col 1 │ Row 1 col 2 │
            ┼─────────────┼─────────────┼
            │ Row 2 col 1 │ Row 2 col 2 │
            ┼─────────────┼─────────────┼

        - ``plain=False, semiplain=True``

          .. code-block:: none

            ─────────────────────────────
              Row 1 col 1
            ─────────────────────────────
                            Row 1 col 2
            ─────────────────────────────
              Row 2 col 1
            ─────────────────────────────
                            Row 2 col 2
            ─────────────────────────────

        - ``plain=True``

          .. code-block:: none

            ╔═════════════════════════════════════════════════════════════════╗
            Row 1 col 1
            ───────────────────────────────────────────────────────────────────
            Row 1 col 2
            ═══════════════════════════════════════════════════════════════════
            Row 2 col 1
            ───────────────────────────────────────────────────────────────────
            Row 2 col 2
            ╚═════════════════════════════════════════════════════════════════╝

        The plain format is probably better, in general, for NLP, and is
        definitely clearer with nested tables (for which the word-wrapping
        algorithm is imperfect). We avoid "heavy" box drawing as it has a
        higher chance of being mangled under Windows.

        """
        if plain and semiplain:
            log.warning("You specified both plain and semiplain; using plain")
            semiplain = False
        middlewidth = width - 2 if width > 2 else 77
        # double
        if plain_table_start is None:
            plain_table_start = "╔" + ("═" * middlewidth) + "╗"
        if plain_table_end is None:
            plain_table_end = "╚" + ("═" * middlewidth) + "╝"
        # heavy
        if plain_table_row_boundary is None:
            plain_table_row_boundary = "═" * (middlewidth + 2)
        # light
        if plain_table_col_boundary is None:
            plain_table_col_boundary = "─" * (middlewidth + 2)

        self.encoding = encoding
        self.width = width
        self.min_col_width = min_col_width
        self.plain = plain
        self.semiplain = semiplain
        self.docx_in_order = docx_in_order
        self.horizontal_char = horizontal_char
        self.vertical_char = vertical_char
        self.junction_char = junction_char
        self.plain_table_start = plain_table_start
        self.plain_table_end = plain_table_end
        self.plain_table_col_boundary = plain_table_col_boundary
        self.plain_table_row_boundary = plain_table_row_boundary
        self.rstrip = rstrip


_DEFAULT_CONFIG = TextProcessingConfig()


# =============================================================================
# Support functions
# =============================================================================


def get_filelikeobject(filename: str = None, blob: bytes = None) -> BinaryIO:
    """
    Open a file-like object.

    Guard the use of this function with ``with``.

    Args:
        filename: for specifying via a filename
        blob: for specifying via an in-memory ``bytes`` object

    Returns:
        a :class:`BinaryIO` object
    """
    if not filename and not blob:
        raise ValueError("no filename and no blob")
    if filename and blob:
        raise ValueError("specify either filename or blob")
    if filename:
        return open(filename, "rb")
    else:
        return io.BytesIO(blob)


# noinspection PyUnusedLocal
def get_file_contents(filename: str = None, blob: bytes = None) -> bytes:
    """
    Returns the binary contents of a file, or of a BLOB.
    """
    if not filename and not blob:
        raise ValueError("no filename and no blob")
    if filename and blob:
        raise ValueError("specify either filename or blob")
    if blob:
        return blob
    with open(filename, "rb") as f:
        return f.read()


def get_chardet_encoding(binary_contents: bytes) -> Optional[str]:
    """
    Guess the character set encoding of the specified ``binary_contents``.
    """
    if not binary_contents:
        return None
    if chardet is None or UniversalDetector is None:
        log.warning("chardet not installed; limits detection of encodings")
        return None
    # METHOD 1
    # http://chardet.readthedocs.io/en/latest/
    #
    # guess = chardet.detect(binary_contents)
    #
    # METHOD 2: faster with large files
    # http://chardet.readthedocs.io/en/latest/
    # https://stackoverflow.com/questions/13857856/split-byte-string-into-lines

    # noinspection PyCallingNonCallable
    detector = UniversalDetector()
    for byte_line in binary_contents.split(b"\n"):
        detector.feed(byte_line)
        if detector.done:
            break
    guess = detector.result
    # Handle result
    if "encoding" not in guess:
        log.warning("Something went wrong within chardet; no encoding")
        return None
    return guess["encoding"]


def get_file_contents_text(
    filename: str = None,
    blob: bytes = None,
    config: TextProcessingConfig = _DEFAULT_CONFIG,
) -> str:
    """
    Returns the string contents of a file, or of a BLOB.
    """
    binary_contents = get_file_contents(filename=filename, blob=blob)
    # 1. Try the encoding the user specified
    if config.encoding:
        try:
            return binary_contents.decode(config.encoding)
        except ValueError:  # of which UnicodeDecodeError is more specific
            # ... https://docs.python.org/3/library/codecs.html
            pass
    # 2. Try the system encoding
    sysdef = sys.getdefaultencoding()
    if sysdef != config.encoding:
        try:
            return binary_contents.decode(sysdef)
        except ValueError:
            pass
    # 3. Try the best guess from chardet
    #    http://chardet.readthedocs.io/en/latest/usage.html
    if chardet:
        guess = chardet.detect(binary_contents)
        if guess["encoding"]:
            return binary_contents.decode(guess["encoding"])
    raise ValueError(
        "Unknown encoding ({})".format(
            f"filename={filename!r}" if filename else "blob"
        )
    )


def get_cmd_output(*args, encoding: str = SYS_ENCODING) -> str:
    """
    Returns text output of a command.
    """
    log.debug("get_cmd_output(): args = {!r}", args)
    p = subprocess.Popen(args, stdout=subprocess.PIPE)
    stdout, stderr = p.communicate()
    return stdout.decode(encoding, errors="ignore")


def get_cmd_output_from_stdin(
    stdint_content_binary: bytes, *args, encoding: str = SYS_ENCODING
) -> str:
    """
    Returns text output of a command, passing binary data in via stdin.
    """
    p = subprocess.Popen(args, stdin=subprocess.PIPE, stdout=subprocess.PIPE)
    stdout, stderr = p.communicate(input=stdint_content_binary)
    return stdout.decode(encoding, errors="ignore")


def rstrip_all_lines(text: str) -> str:
    """
    Right-strips all lines in a string and returns the result.
    """
    return "\n".join(line.rstrip() for line in text.splitlines())


# =============================================================================
# PDF
# =============================================================================

# noinspection PyUnresolvedReferences,PyUnusedLocal
def convert_pdf_to_txt(
    filename: str = None,
    blob: bytes = None,
    config: TextProcessingConfig = _DEFAULT_CONFIG,
) -> str:
    """
    Converts a PDF file to text.
    Pass either a filename or a binary object.
    """
    pdftotext = tools["pdftotext"]
    if pdftotext:  # External command method
        if filename:
            return get_cmd_output(pdftotext, filename, "-")
        else:
            return get_cmd_output_from_stdin(blob, pdftotext, "-", "-")
    elif pdfminer:  # Memory-hogging method
        with get_filelikeobject(filename, blob) as fp:
            rsrcmgr = pdfminer.pdfinterp.PDFResourceManager()
            retstr = StringIO()
            codec = ENCODING
            laparams = pdfminer.layout.LAParams()
            device = pdfminer.converter.TextConverter(
                rsrcmgr, retstr, codec=codec, laparams=laparams
            )
            interpreter = pdfminer.pdfinterp.PDFPageInterpreter(
                rsrcmgr, device
            )
            password = ""
            maxpages = 0
            caching = True
            pagenos = set()
            for page in pdfminer.pdfpage.PDFPage.get_pages(
                fp,
                pagenos,
                maxpages=maxpages,
                password=password,
                caching=caching,
                check_extractable=True,
            ):
                interpreter.process_page(page)
            text = retstr.getvalue().decode(ENCODING)
        return text
    else:
        raise AssertionError("No PDF-reading tool available")


def availability_pdf() -> bool:
    """
    Is a PDF-to-text tool available?
    """
    pdftotext = tools["pdftotext"]
    if pdftotext:
        return True
    elif pdfminer:
        log.warning(
            "PDF conversion: pdftotext missing; "
            "using pdfminer (less efficient)"
        )
        return True
    else:
        return False


# =============================================================================
# DOCX
# =============================================================================

# -----------------------------------------------------------------------------
# In a D.I.Y. fashion
# -----------------------------------------------------------------------------
# DOCX specification: http://www.ecma-international.org/news/TC45_current_work/TC45_available_docs.htm  # noqa

DOCX_HEADER_FILE_REGEX = re.compile("word/header[0-9]*.xml")
DOCX_DOC_FILE = "word/document.xml"
DOCX_FOOTER_FILE_REGEX = re.compile("word/footer[0-9]*.xml")
DOCX_SCHEMA_URL = (
    "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
)


def docx_qn(tagroot):
    return f"{{{DOCX_SCHEMA_URL}}}{tagroot}"


DOCX_TEXT = docx_qn("t")
DOCX_TABLE = docx_qn(
    "tbl"
)  # https://github.com/python-openxml/python-docx/blob/master/docx/table.py  # noqa
DOCX_TAB = docx_qn("tab")
DOCX_NEWLINES = [docx_qn("br"), docx_qn("cr")]
DOCX_NEWPARA = docx_qn("p")
DOCX_TABLE_ROW = docx_qn("tr")
DOCX_TABLE_CELL = docx_qn("tc")


def gen_xml_files_from_docx(fp: BinaryIO) -> Iterator[str]:
    """
    Generate XML files (as strings) from a DOCX file.

    Args:
        fp: :class:`BinaryIO` object for reading the ``.DOCX`` file

    Yields:
        the string contents of each individual XML file within the ``.DOCX``
        file

    Raises:
        zipfile.BadZipFile: if the zip is unreadable (encrypted?)

    """
    try:
        z = zipfile.ZipFile(fp)
        filelist = z.namelist()
        for filename in filelist:
            if DOCX_HEADER_FILE_REGEX.match(filename):
                yield z.read(filename).decode("utf8")
        yield z.read(DOCX_DOC_FILE)
        for filename in filelist:
            if DOCX_FOOTER_FILE_REGEX.match(filename):
                yield z.read(filename).decode("utf8")
    except zipfile.BadZipFile:
        # Clarify the error:
        raise zipfile.BadZipFile("File is not a zip file - encrypted DOCX?")


class DocxFragment(object):
    """
    Representation of a line, or multiple lines, which may or may not need
    word-wrapping.
    """

    # noinspection PyShadowingNames
    def __init__(self, text: str, wordwrap: bool = True) -> None:
        self.text = text
        self.wordwrap = wordwrap


def docx_gen_wordwrapped_fragments(
    fragments: Iterable[DocxFragment], width: int
) -> Generator[str, None, None]:
    """
    Generates word-wrapped fragments.
    """
    to_wrap = []  # type: List[DocxFragment]

    def yield_wrapped():
        """
        Yield the word-wrapped stuff to date.
        """
        nonlocal to_wrap
        if to_wrap:
            block = "".join(x.text for x in to_wrap)
            wrapped = "\n".join(
                wordwrap(line, width) for line in block.splitlines()
            )
            yield wrapped
            to_wrap.clear()

    for f in fragments:
        if f.wordwrap:
            # Add it to the current wrapping block.
            to_wrap.append(f)
        else:
            # Yield the wrapped stuff to date
            yield from yield_wrapped()
            # Yield the new, unwrapped
            yield f.text

    yield from yield_wrapped()  # any leftovers


def docx_wordwrap_fragments(
    fragments: Iterable[DocxFragment], width: int
) -> str:
    """
    Joins multiple fragments and word-wraps them as necessary.
    """
    return "".join(docx_gen_wordwrapped_fragments(fragments, width))


def docx_gen_fragments_from_xml_node(
    node: ElementTree.Element, level: int, config: TextProcessingConfig
) -> Generator[DocxFragment, None, None]:
    """
    Returns text from an XML node within a DOCX file.

    Args:
        node: an XML node
        level: current level in XML hierarchy (used for recursion; start level
            is 0)
        config: :class:`TextProcessingConfig` control object

    Returns:
        contents as a string

    """
    tag = node.tag  # for speed
    log.debug("Level {}, tag {}", level, tag)
    if tag == DOCX_TEXT:
        log.debug("Text: {!r}", node.text)
        yield DocxFragment(node.text or "")
    elif tag == DOCX_TAB:
        log.debug("Tab")
        yield DocxFragment("\t")
    elif tag in DOCX_NEWLINES:  # rarely used? Mostly "new paragraph"
        log.debug("Newline")
        yield DocxFragment("\n")
    elif tag == DOCX_NEWPARA:  # Note that e.g. all table cells start with this
        log.debug("New paragraph")
        yield DocxFragment("\n\n")
        # One or two newlines? Clarity better with two -- word-wrapping means
        # that "single" source lines can take up multiple lines in text format.
        # So we need a gap between lines to ensure paragraph separation is
        # visible -- i.e. two newlines.

    if tag == DOCX_TABLE:
        log.debug("Table")
        yield DocxFragment("\n", wordwrap=False)
        yield DocxFragment(
            docx_table_from_xml_node(node, level, config), wordwrap=False
        )
    else:
        for child in node:
            for fragment in docx_gen_fragments_from_xml_node(
                child, level + 1, config
            ):
                yield fragment


def docx_text_from_xml_node(
    node: ElementTree.Element, level: int, config: TextProcessingConfig
) -> str:
    """
    Returns text from an XML node within a DOCX file.

    Args:
        node: an XML node
        level: current level in XML hierarchy (used for recursion; start level
            is 0)
        config: :class:`TextProcessingConfig` control object

    Returns:
        contents as a string

    """
    return docx_wordwrap_fragments(
        docx_gen_fragments_from_xml_node(node, level, config), config.width
    )


def docx_text_from_xml(xml: str, config: TextProcessingConfig) -> str:
    """
    Converts an XML tree of a DOCX file to string contents.

    Args:
        xml: raw XML text
        config: :class:`TextProcessingConfig` control object

    Returns:
        contents as a string
    """
    root = ElementTree.fromstring(xml)
    return docx_text_from_xml_node(root, 0, config)


class CustomDocxParagraph(object):
    """
    Represents a paragraph of text in a DOCX file.
    """

    def __init__(self, text: str = "") -> None:
        self.text = text or ""

    def __repr__(self) -> str:
        return f"CustomDocxParagraph(text={self.text!r})"


class CustomDocxTableCell(object):
    """
    Represents a cell within a table of a DOCX file.
    May contain several paragraphs.
    """

    def __init__(self, paragraphs: List[CustomDocxParagraph] = None) -> None:
        self.paragraphs = paragraphs or []

    def add_paragraph(self, text: str) -> None:
        self.paragraphs.append(CustomDocxParagraph(text))

    def __repr__(self) -> str:
        return f"CustomDocxTableCell(paragraphs={self.paragraphs!r})"


class CustomDocxTableRow(object):
    """
    Represents a row within a table of a DOCX file.
    May contain several cells (one per column).
    """

    def __init__(self, cells: List[CustomDocxTableCell] = None) -> None:
        self.cells = cells or []

    def add_cell(self, cell: CustomDocxTableCell) -> None:
        self.cells.append(cell)

    def new_cell(self) -> None:
        self.cells.append(CustomDocxTableCell())

    def add_paragraph(self, text: str) -> None:
        self.cells[-1].add_paragraph(text)

    def __repr__(self) -> str:
        return f"CustomDocxTableRow(cells={self.cells!r})"


class CustomDocxTable(object):
    """
    Represents a table of a DOCX file.
    May contain several rows.
    """

    def __init__(self, rows: List[CustomDocxTableRow] = None) -> None:
        self.rows = rows or []

    def add_row(self, row: CustomDocxTableRow) -> None:
        self.rows.append(row)

    def new_row(self) -> None:
        self.rows.append(CustomDocxTableRow())

    def new_cell(self) -> None:
        self.rows[-1].new_cell()

    def add_paragraph(self, text: str) -> None:
        self.rows[-1].add_paragraph(text)

    def __repr__(self) -> str:
        return f"CustomDocxTable(rows={self.rows!r})"


def docx_table_from_xml_node(
    table_node: ElementTree.Element, level: int, config: TextProcessingConfig
) -> str:
    """
    Converts an XML node representing a DOCX table into a textual
    representation.

    Args:
        table_node: XML node
        level: current level in XML hierarchy (used for recursion; start level
            is 0)
        config: :class:`TextProcessingConfig` control object

    Returns:
        string representation

    """
    table = CustomDocxTable()
    for row_node in table_node:
        if row_node.tag != DOCX_TABLE_ROW:
            continue
        table.new_row()
        for cell_node in row_node:
            if cell_node.tag != DOCX_TABLE_CELL:
                continue
            table.new_cell()
            for para_node in cell_node:
                text = docx_text_from_xml_node(para_node, level, config)
                if text:
                    table.add_paragraph(text)
    return docx_process_table(table, config)


# -----------------------------------------------------------------------------
# Generic
# -----------------------------------------------------------------------------


def wordwrap(text: str, width: int) -> str:
    """
    Word-wraps text.

    Args:
        text:
            text to process (will be treated as a single line)
        width:
            width to word-wrap to (or 0 to skip word wrapping)

    Returns:
        wrapped text


    .. code-block:: python

        from cardinal_pythonlib.extract_text import *
        text = "Here is a very long line that may be word-wrapped. " * 50
        print(docx_wordwrap(text, 80))
    """
    if not text:
        return ""
    if width:
        return "\n".join(textwrap.wrap(text, width=width))
    return text


def docx_process_table(
    table: DOCX_TABLE_TYPE, config: TextProcessingConfig
) -> str:
    """
    Converts a DOCX table to text.

    Structure representing a DOCX table:

    .. code-block:: none

        table
            .rows[]
                .cells[]
                    .paragraphs[]
                        .text

    That's the structure of a :class:`docx.table.Table` object, but also of our
    homebrew creation, :class:`CustomDocxTable`.

    - The ``plain`` and ``semiplain`` options are implemented via the
      :class:`TextProcessingConfig`.

    - Note also that the grids in DOCX files can have varying number of cells
      per row, e.g.

      .. code-block:: none

            +---+---+---+
            | 1 | 2 | 3 |
            +---+---+---+
            | 1 | 2 |
            +---+---+

    """

    def get_cell_text(cell_) -> str:
        cellparagraphs = [
            paragraph.text.strip() for paragraph in cell_.paragraphs
        ]
        cellparagraphs = [x for x in cellparagraphs if x]
        return "\n\n".join(cellparagraphs)

    if config.plain:
        # ---------------------------------------------------------------------
        # Plain -- good for NLP and better for word-wrapping
        # ---------------------------------------------------------------------
        lines = [config.plain_table_start]  # type: List[str]
        for r, row in enumerate(table.rows):
            if r > 0:
                lines.append(config.plain_table_row_boundary)
            for c, cell in enumerate(row.cells):
                if c > 0:
                    lines.append(config.plain_table_col_boundary)
                lines.append(get_cell_text(cell))
        lines.append(config.plain_table_end)
        return "\n".join(lines)
    else:
        # ---------------------------------------------------------------------
        # Full table visualization, or semiplain
        # ---------------------------------------------------------------------
        ncols = 1
        # noinspection PyTypeChecker
        for row in table.rows:
            ncols = max(ncols, len(row.cells))
        pt = prettytable.PrettyTable(
            field_names=list(range(ncols)),
            encoding=ENCODING,
            header=False,
            border=True,
            hrules=prettytable.ALL,
            vrules=prettytable.NONE if config.semiplain else prettytable.ALL,
            # Can we use UTF-8 special characters?
            # Even under Windows, sys.getdefaultencoding() returns "utf-8"
            # (under Python 3.6.8, Windows 6.1.7601 = Windows Server 2008 R2).
            # The advantage would be that these characters are not likely to
            # influence any form of NLP.
            horizontal_char=config.horizontal_char,  # default "-"
            vertical_char=config.vertical_char,  # default "|"
            junction_char=config.junction_char,  # default "+"
        )
        pt.align = "l"
        pt.valign = "t"
        pt.max_width = max(config.width // ncols, config.min_col_width)
        if config.semiplain:
            # noinspection PyTypeChecker
            for row in table.rows:
                for i, cell in enumerate(row.cells):
                    n_before = i
                    n_after = ncols - i - 1
                    # ... use ncols, not len(row.cells), since "cells per row"
                    #     is not constant, but prettytable wants a fixed
                    #     number. (changed in v0.2.8)
                    ptrow = (
                        [""] * n_before
                        + [get_cell_text(cell)]
                        + [""] * n_after
                    )
                    assert len(ptrow) == ncols
                    pt.add_row(ptrow)
        else:
            # noinspection PyTypeChecker
            for row in table.rows:
                ptrow = []  # type: List[str]
                # noinspection PyTypeChecker
                for cell in row.cells:
                    ptrow.append(get_cell_text(cell))
                ptrow += [""] * (ncols - len(ptrow))  # added in v0.2.8
                assert len(ptrow) == ncols
                pt.add_row(ptrow)
        return pt.get_string()


# -----------------------------------------------------------------------------
# With the docx library
# -----------------------------------------------------------------------------

_ = '''
# noinspection PyProtectedMember,PyUnresolvedReferences
def docx_docx_iter_block_items(parent: DOCX_CONTAINER_TYPE) \
        -> Iterator[DOCX_BLOCK_ITEM_TYPE]:
    """
    Iterate through items of a DOCX file.

    See https://github.com/python-openxml/python-docx/issues/40.

    Yield each paragraph and table child within ``parent``, in document order.
    Each returned value is an instance of either :class:`Table` or
    :class:`Paragraph`. ``parent`` would most commonly be a reference to a main
    :class:`Document` object, but also works for a :class:`_Cell` object, which
    itself can contain paragraphs and tables.

    NOTE: uses internals of the ``python-docx`` (``docx``) library; subject to
    change; this version works with ``docx==0.8.5``.
    """
    if isinstance(parent, docx.document.Document):
        parent_elm = parent.element.body
    elif isinstance(parent, docx.table._Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("something's not right")

    for child in parent_elm.iterchildren():
        if isinstance(child, docx.oxml.text.paragraph.CT_P):
            yield docx.text.paragraph.Paragraph(child, parent)
        elif isinstance(child, docx.oxml.table.CT_Tbl):
            yield docx.table.Table(child, parent)


# noinspection PyUnresolvedReferences
def docx_docx_gen_text(doc: DOCX_DOCUMENT_TYPE,
                       config: TextProcessingConfig) -> Iterator[str]:
    """
    Iterate through a DOCX file and yield text.

    Args:
        doc: DOCX document to process
        config: :class:`TextProcessingConfig` control object

    Yields:
        pieces of text (paragraphs)

    """
    if in_order:
        for thing in docx_docx_iter_block_items(doc):
            if isinstance(thing, docx.text.paragraph.Paragraph):
                yield docx_process_simple_text(thing.text, config.width)
            elif isinstance(thing, docx.table.Table):
                yield docx_process_table(thing, config)
    else:
        for paragraph in doc.paragraphs:
            yield docx_process_simple_text(paragraph.text, config.width)
        for table in doc.tables:
            yield docx_process_table(table, config)
'''


# noinspection PyUnusedLocal
def convert_docx_to_text(
    filename: str = None,
    blob: bytes = None,
    config: TextProcessingConfig = _DEFAULT_CONFIG,
) -> str:
    """
    Converts a DOCX file to text.
    Pass either a filename or a binary object.

    Args:
        filename: filename to process
        blob: binary ``bytes`` object to process
        config: :class:`TextProcessingConfig` control object

    Returns:
        text contents

    Notes:

    - Old ``docx`` (https://pypi.python.org/pypi/python-docx) has been
      superseded (see https://github.com/mikemaccana/python-docx).

      - ``docx.opendocx(file)`` uses :class:`zipfile.ZipFile`, which can take
        either a filename or a file-like object
        (https://docs.python.org/2/library/zipfile.html).

      - Method was:

        .. code-block:: python

            with get_filelikeobject(filename, blob) as fp:
                document = docx.opendocx(fp)
                paratextlist = docx.getdocumenttext(document)
            return '\n\n'.join(paratextlist)

    - Newer ``docx`` is python-docx

      - https://pypi.python.org/pypi/python-docx
      - https://python-docx.readthedocs.org/en/latest/
      - https://stackoverflow.com/questions/25228106

      However, it uses ``lxml``, which has C dependencies, so it doesn't always
      install properly on e.g. bare Windows machines.

      PERFORMANCE of my method:

      - nice table formatting
      - but tables grouped at end, not in sensible places
      - can iterate via ``doc.paragraphs`` and ``doc.tables`` but not in
        true document order, it seems
      - others have noted this too:

        - https://github.com/python-openxml/python-docx/issues/40
        - https://github.com/deanmalmgren/textract/pull/92

    - ``docx2txt`` is at https://pypi.python.org/pypi/docx2txt/0.6; this is
      pure Python. Its command-line function appears to be for Python 2 only
      (2016-04-21: crashes under Python 3; is due to an encoding bug). However,
      it seems fine as a library. It doesn't handle in-memory blobs properly,
      though, so we need to extend it.

      PERFORMANCE OF ITS ``process()`` function:

      - all text comes out
      - table text is in a sensible place
      - table formatting is lost.

    - Other manual methods (not yet implemented):
      https://etienned.github.io/posts/extract-text-from-word-docx-simply/.

      Looks like it won't deal with header stuff (etc.) that ``docx2txt``
      handles.

    - Upshot: we need a DIY version.

    - See also this "compile lots of techniques" libraries, which has C
      dependencies: https://textract.readthedocs.org/en/latest/

    """

    text = ""
    with get_filelikeobject(filename, blob) as fp:
        for xml in gen_xml_files_from_docx(fp):
            text += docx_text_from_xml(xml, config)
    return text

    # elif docx:
    #     with get_filelikeobject(filename, blob) as fp:
    #         # noinspection PyUnresolvedReferences
    #         document = docx.Document(fp)
    #         return '\n\n'.join(
    #             docx_docx_gen_text(document, config))
    # elif docx2txt:
    #     if filename:
    #         return docx2txt.process(filename)
    #     else:
    #         raise NotImplementedError("docx2txt BLOB handling not written")
    # else:
    #     raise AssertionError("No DOCX-reading tool available")


# =============================================================================
# ODT
# =============================================================================

# noinspection PyUnusedLocal
def convert_odt_to_text(
    filename: str = None,
    blob: bytes = None,
    config: TextProcessingConfig = _DEFAULT_CONFIG,
) -> str:
    """
    Converts an OpenOffice ODT file to text.

    Pass either a filename or a binary object.
    """
    # We can't use exactly the same method as for DOCX files, using docx:
    # sometimes that works, but sometimes it falls over with:
    # KeyError: "There is no item named 'word/document.xml' in the archive"
    with get_filelikeobject(filename, blob) as fp:
        z = zipfile.ZipFile(fp)
        tree = ElementTree.fromstring(z.read("content.xml"))
        # ... may raise zipfile.BadZipfile
        textlist = []  # type: List[str]
        for element in tree.iter():
            if element.text:
                textlist.append(element.text.strip())
    return "\n\n".join(textlist)


# =============================================================================
# HTML
# =============================================================================

# noinspection PyUnusedLocal
def convert_html_to_text(
    filename: str = None,
    blob: bytes = None,
    config: TextProcessingConfig = _DEFAULT_CONFIG,
) -> str:
    """
    Converts HTML to text.
    """
    with get_filelikeobject(filename, blob) as fp:
        soup = bs4.BeautifulSoup(fp)
        return soup.get_text()


# =============================================================================
# XML
# =============================================================================

# noinspection PyUnusedLocal
def convert_xml_to_text(
    filename: str = None,
    blob: bytes = None,
    config: TextProcessingConfig = _DEFAULT_CONFIG,
) -> str:
    """
    Converts XML to text.
    """
    with get_filelikeobject(filename, blob) as fp:
        soup = bs4.BeautifulStoneSoup(fp)
        return soup.get_text()


# =============================================================================
# RTF
# =============================================================================

# noinspection PyUnresolvedReferences,PyUnusedLocal
def convert_rtf_to_text(
    filename: str = None,
    blob: bytes = None,
    config: TextProcessingConfig = _DEFAULT_CONFIG,
) -> str:
    """
    Converts RTF to text.
    """
    unrtf = tools["unrtf"]
    if unrtf:  # Best
        args = [unrtf, "--text", "--nopict"]
        if UNRTF_SUPPORTS_QUIET:
            args.append("--quiet")
        if filename:
            args.append(filename)
            return get_cmd_output(*args)
        else:
            return get_cmd_output_from_stdin(blob, *args)
    elif pyth:  # Very memory-consuming:
        # https://github.com/brendonh/pyth/blob/master/pyth/plugins/rtf15/reader.py  # noqa
        with get_filelikeobject(filename, blob) as fp:
            doc = pyth.plugins.rtf15.reader.Rtf15Reader.read(fp)
        return pyth.plugins.plaintext.writer.PlaintextWriter.write(
            doc
        ).getvalue()
    else:
        raise AssertionError("No RTF-reading tool available")


def availability_rtf() -> bool:
    """
    Is an RTF processor available?
    """
    unrtf = tools["unrtf"]
    if unrtf:
        return True
    elif pyth:
        log.warning(
            "RTF conversion: unrtf missing; " "using pyth (less efficient)"
        )
        return True
    else:
        return False


# =============================================================================
# DOC
# =============================================================================

# noinspection PyUnusedLocal
def convert_doc_to_text(
    filename: str = None,
    blob: bytes = None,
    config: TextProcessingConfig = _DEFAULT_CONFIG,
) -> str:
    """
    Converts Microsoft Word DOC files to text.
    """
    antiword = tools["antiword"]
    if antiword:
        if filename:
            return get_cmd_output(antiword, "-w", str(config.width), filename)
        else:
            return get_cmd_output_from_stdin(
                blob, antiword, "-w", str(config.width), "-"
            )
    else:
        raise AssertionError("No DOC-reading tool available")


def availability_doc() -> bool:
    """
    Is a DOC processor available?
    """
    antiword = tools["antiword"]
    return bool(antiword)


# =============================================================================
# Anything
# =============================================================================

# noinspection PyUnusedLocal
def convert_anything_to_text(
    filename: str = None,
    blob: bytes = None,
    config: TextProcessingConfig = _DEFAULT_CONFIG,
) -> str:
    """
    Convert arbitrary files to text, using ``strings`` or ``strings2``.
    (``strings`` is a standard Unix command to get text from any old rubbish.)
    """
    strings = tools["strings"] or tools["strings2"]
    if strings:
        if filename:
            return get_cmd_output(strings, filename)
        else:
            return get_cmd_output_from_stdin(blob, strings)
    else:
        raise AssertionError("No fallback string-reading tool available")


def availability_anything() -> bool:
    """
    Is a generic "something-to-text" processor available?
    """

    strings = tools["strings"] or tools["strings2"]
    return bool(strings)


# =============================================================================
# Decider
# =============================================================================

ext_map = {
    # Converter functions must be of the form: func(filename, blob, config).
    # Availability must be either a boolean literal or a function that takes no
    # params.
    ".csv": {CONVERTER: get_file_contents_text, AVAILABILITY: True},
    ".doc": {CONVERTER: convert_doc_to_text, AVAILABILITY: availability_doc},
    ".docm": {CONVERTER: convert_docx_to_text, AVAILABILITY: True},
    ".docx": {CONVERTER: convert_docx_to_text, AVAILABILITY: True},
    ".dot": {CONVERTER: convert_doc_to_text, AVAILABILITY: availability_doc},
    ".htm": {CONVERTER: convert_html_to_text, AVAILABILITY: True},
    ".html": {CONVERTER: convert_html_to_text, AVAILABILITY: True},
    ".log": {CONVERTER: get_file_contents_text, AVAILABILITY: True},
    # .msg is often Outlook binary, not text
    #
    # '.msg': {
    #     CONVERTER: get_file_contents_text,
    #     AVAILABILITY: True,
    # },
    ".odt": {CONVERTER: convert_odt_to_text, AVAILABILITY: True},
    ".pdf": {CONVERTER: convert_pdf_to_txt, AVAILABILITY: availability_pdf},
    ".rtf": {CONVERTER: convert_rtf_to_text, AVAILABILITY: availability_rtf},
    ".txt": {CONVERTER: get_file_contents_text, AVAILABILITY: True},
    ".xml": {CONVERTER: convert_xml_to_text, AVAILABILITY: True},
    None: {  # fallback
        CONVERTER: convert_anything_to_text,
        AVAILABILITY: availability_anything,
    },
}


def document_to_text(
    filename: str = None,
    blob: bytes = None,
    extension: str = None,
    config: TextProcessingConfig = _DEFAULT_CONFIG,
) -> str:
    """
    Converts a document to text.

    This function selects a processor based on the file extension (either from
    the filename, or, in the case of a BLOB, the extension specified manually
    via the ``extension`` parameter).

    Pass either a filename or a binary object.

    Args:

        filename:
            the filename to read
        blob:
            binary content (alternative to ``filename``)
        extension:
            file extension, used as a hint when ``blob`` is used
        config:
            an optional :class:`TextProcessingConfig` object

    Returns:

        Returns a string if the file was processed (potentially an empty
        string).

    Raises:

        Raises an exception for malformed arguments, missing files, bad
        filetypes, etc.
    """
    if not filename and not blob:
        raise ValueError("document_to_text: no filename and no blob")
    if filename and blob:
        raise ValueError("document_to_text: specify either filename or blob")
    if blob and not extension:
        raise ValueError("document_to_text: need extension hint for blob")
    if filename:
        stub, extension = os.path.splitext(filename)
    else:
        if extension[0] != ".":
            extension = "." + extension
    extension = extension.lower()

    # Ensure blob is an appropriate type
    log.debug(
        f"filename: {filename}, blob type: {type(blob)}, "
        f"blob length: {len(blob) if blob is not None else None}, "
        f"extension: {extension}"
    )

    # If we were given a filename and the file doesn't exist, don't bother.
    if filename and not os.path.isfile(filename):
        raise ValueError(f"document_to_text: no such file: {filename!r}")

    # Choose method
    info = ext_map.get(extension)
    if info is None:
        log.warning("Unknown filetype: {}; using generic tool", extension)
        info = ext_map[None]
    func = info[CONVERTER]
    text = func(filename, blob, config)
    if config.rstrip:
        text = rstrip_all_lines(text)
    return text


def is_text_extractor_available(extension: str) -> bool:
    """
    Is a text extractor available for the specified extension?
    """
    if extension is not None:
        extension = extension.lower()
    info = ext_map.get(extension)
    if info is None:
        return False
    availability = info[AVAILABILITY]
    if type(availability) == bool:
        return availability
    elif callable(availability):
        return availability()
    else:
        raise ValueError(f"Bad information object for extension: {extension}")


def require_text_extractor(extension: str) -> None:
    """
    Require that a text extractor is available for the specified extension,
    or raise :exc:`ValueError`.
    """
    if not is_text_extractor_available(extension):
        raise ValueError(
            f"No text extractor available for extension: {extension}"
        )


# =============================================================================
# main, for command-line use
# =============================================================================


def main() -> None:
    """
    Command-line processor. See ``--help`` for details.
    """
    logging.basicConfig(level=logging.DEBUG)
    parser = argparse.ArgumentParser(
        formatter_class=argparse.ArgumentDefaultsHelpFormatter
    )
    parser.add_argument("inputfile", nargs="?", help="Input file name")
    parser.add_argument(
        "--availability",
        nargs="*",
        help="File extensions to check availability for (use a '.' prefix, "
        "and use the special extension 'None' to check the fallback "
        "processor",
    )
    parser.add_argument(
        "--plain",
        action="store_true",
        help="Use plainest format (re e.g. table layouts)",
    )
    parser.add_argument(
        "--semiplain",
        action="store_true",
        help="Use semi-plain format (re e.g. table layouts)",
    )
    parser.add_argument(
        "--width", type=int, default=DEFAULT_WIDTH, help="Word wrapping width"
    )
    parser.add_argument(
        "--min-col-width",
        type=int,
        default=DEFAULT_MIN_COL_WIDTH,
        help="Minimum column width for tables",
    )
    args = parser.parse_args()
    if args.availability:
        for ext in args.availability:
            if ext.lower() == "none":
                ext = None
            available = is_text_extractor_available(ext)
            print(f"Extractor for extension {ext} present: {available}")
        return
    if not args.inputfile:
        parser.print_help(sys.stderr)
        return
    config = TextProcessingConfig(
        width=args.width,
        min_col_width=args.min_col_width,
        plain=args.plain,
        semiplain=args.semiplain,
    )
    result = document_to_text(filename=args.inputfile, config=config)
    if result is None:
        return
    else:
        print(result)


if __name__ == "__main__":
    main()


# *** antiword -w width
